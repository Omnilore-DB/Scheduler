#!/usr/bin/env python3
"""Render Project_Docs_D7 Markdown sources to simple DOCX and PDF files.

This script is intentionally dependency-light so the handoff docs can be
refreshed from the checked-in Markdown sources in local Codex/Claude sessions.
Markdown remains the source of truth; DOCX/PDF files are rendered deliverables.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    ListFlowable,
    ListItem,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


DOCS = [
    "System_Overview_and_Architecture.md",
    "Setup_and_Deployment_Guide.md",
    "Operations_Runbook.md",
    "API_and_Data_Reference.md",
    "Testing_and_QA_Summary.md",
    "Security_Privacy_Accessibility_UX_Notes.md",
    "Backlog_Known_Issues_Roadmap.md",
    "Demo_Video_Script_and_Checklist.md",
    "Handoff_Checklist_and_Verification_Log.md",
    "Handoff_Package_Manifest.md",
    "Stakeholder_Email_Cover_Note.md",
]


def inline_markup(text: str) -> str:
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"`([^`]+)`", r"<font face='Courier'>\1</font>", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
    return text


def strip_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text


def parse_markdown(path: Path):
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue

        image_match = re.match(r"!\[[^\]]*\]\(([^)]+)\)", line.strip())
        if image_match:
            yield ("image", image_match.group(1))
            i += 1
            continue

        if line.startswith("```"):
            lang = line.strip("`").strip()
            block = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            yield ("code", lang, "\n".join(block))
            continue

        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = []
            for row in table_lines:
                cells = [strip_inline(cell.strip()) for cell in row.strip().strip("|").split("|")]
                if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    continue
                rows.append(cells)
            if rows:
                yield ("table", rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            yield ("heading", len(heading.group(1)), strip_inline(heading.group(2)))
            i += 1
            continue

        if re.match(r"^\s*[-*]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                items.append(strip_inline(re.sub(r"^\s*[-*]\s+", "", lines[i]).strip()))
                i += 1
            yield ("bullets", items)
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                items.append(strip_inline(re.sub(r"^\s*\d+\.\s+", "", lines[i]).strip()))
                i += 1
            yield ("numbers", items)
            continue

        if line.startswith(">"):
            quote = []
            while i < len(lines) and lines[i].startswith(">"):
                quote.append(strip_inline(lines[i].lstrip("> ").strip()))
                i += 1
            yield ("quote", " ".join(quote))
            continue

        paragraph = [line.strip()]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(
            r"^(#{1,6})\s+|```|!\[[^\]]*\]\(|\||\s*[-*]\s+|\s*\d+\.\s+|>",
            lines[i],
        ):
            paragraph.append(lines[i].strip())
            i += 1
        yield ("paragraph", " ".join(paragraph))


def render_docx(md_path: Path, out_path: Path) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    for item in parse_markdown(md_path):
        kind = item[0]
        if kind == "heading":
            _, level, text = item
            doc.add_heading(text, level=min(level, 4))
        elif kind == "paragraph":
            doc.add_paragraph(strip_inline(item[1]))
        elif kind == "quote":
            p = doc.add_paragraph(strip_inline(item[1]))
            p.style = "Intense Quote"
        elif kind == "bullets":
            for text in item[1]:
                doc.add_paragraph(text, style="List Bullet")
        elif kind == "numbers":
            for text in item[1]:
                doc.add_paragraph(text, style="List Number")
        elif kind == "code":
            _, _lang, code = item
            p = doc.add_paragraph()
            run = p.add_run(code)
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
        elif kind == "table":
            rows = item[1]
            cols = max(len(row) for row in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for r, row in enumerate(rows):
                for c in range(cols):
                    table.cell(r, c).text = row[c] if c < len(row) else ""
        elif kind == "image":
            img = (md_path.parent / item[1]).resolve()
            if img.exists():
                doc.add_picture(str(img), width=Inches(6.0))

    doc.save(out_path)


def render_pdf(md_path: Path, out_path: Path) -> None:
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CodeBlock", fontName="Courier", fontSize=8, leading=10))
    styles.add(ParagraphStyle(name="Quote", parent=styles["BodyText"], leftIndent=18, textColor=colors.HexColor("#555555")))

    story = []
    for item in parse_markdown(md_path):
        kind = item[0]
        if kind == "heading":
            _, level, text = item
            style = "Title" if level == 1 else f"Heading{min(level, 3)}"
            story.append(Paragraph(inline_markup(text), styles[style]))
            story.append(Spacer(1, 0.08 * inch))
        elif kind == "paragraph":
            story.append(Paragraph(inline_markup(item[1]), styles["BodyText"]))
            story.append(Spacer(1, 0.08 * inch))
        elif kind == "quote":
            story.append(Paragraph(inline_markup(item[1]), styles["Quote"]))
            story.append(Spacer(1, 0.08 * inch))
        elif kind in {"bullets", "numbers"}:
            flow_items = [ListItem(Paragraph(inline_markup(text), styles["BodyText"])) for text in item[1]]
            story.append(ListFlowable(flow_items, bulletType="1" if kind == "numbers" else "bullet"))
            story.append(Spacer(1, 0.08 * inch))
        elif kind == "code":
            story.append(Preformatted(item[2], styles["CodeBlock"]))
            story.append(Spacer(1, 0.08 * inch))
        elif kind == "table":
            rows = item[1]
            data = [[Paragraph(inline_markup(cell), styles["BodyText"]) for cell in row] for row in rows]
            table = Table(data, repeatRows=1)
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e9eef5")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#b8c2cc")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.append(table)
            story.append(Spacer(1, 0.12 * inch))
        elif kind == "image":
            img = (md_path.parent / item[1]).resolve()
            if img.exists():
                story.append(Image(str(img), width=6.2 * inch, height=3.5 * inch, kind="proportional"))
                story.append(Spacer(1, 0.12 * inch))

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=letter,
        rightMargin=0.65 * inch,
        leftMargin=0.65 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
    )
    doc.build(story)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docs-dir", default="Project_Docs_D7")
    parser.add_argument("files", nargs="*", help="Markdown files to render. Defaults to D7 PDF/DOCX docs.")
    args = parser.parse_args()

    docs_dir = Path(args.docs_dir)
    files = args.files or DOCS
    for name in files:
        md_path = docs_dir / name
        if not md_path.exists():
            raise FileNotFoundError(md_path)
        stem = md_path.with_suffix("")
        render_docx(md_path, stem.with_suffix(".docx"))
        render_pdf(md_path, stem.with_suffix(".pdf"))
        print(f"rendered {md_path.name}")


if __name__ == "__main__":
    main()
